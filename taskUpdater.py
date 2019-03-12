#! /anaconda3/bin/python

import os 
import sys
import docx

def isFile(file):
	return os.path.isfile(file)

def getFileNames():
	dirs = os.listdir()				# Get all file names in the current directory
	file_names = []

	for file in dirs:
		if isFileExist(file):
			file_name, ext = file.split('.')

			if ext == 'py':
				file_names.append(file_name)

	return file_names

def isFileExist(file_name):
	if os.path.exists(file_name):
		return True
	return False

def createDoc(file_name):
	print("%s file name in createDoc"%file_name)
	if not isFileExist(file_name):
		file, ext = file_name.split('.')
		doc = docx.Document()
		doc.add_heading("%s Script Tasks"%(file.capitalize()))
		doc.save(file_name)
		print("File created.")
	else:
		print("File already exists.")

def printOptions(file_names):
	for num, file in enumerate(file_names):
		print("%d. %s"%(num+1, file))

def chooseFile():
	choice = int(input("Choose file number: "))
	return choice

def addTask(file_name):
	#print("%s name in addTask"%file_name)
	doc = docx.Document(file_name)
	while(True):
		para = input("Add task or enter q to quit:\n")
		print('')
		if para is 'q' or para is 'Q':
			break;
		paragraph = doc.add_paragraph(para)
		paragraph.style = 'List Number'

	doc.save(file_name)
	print("Tasks saved to doc %s"%(file_name.capitalize()))

def createDocFileNames(file_names):
	for i in range(len(file_names)):
		file_names[i] += '.docx'

	return file_names

def main():
	#print("Getting all the file names")
	file_names = getFileNames()

	print("File names:")
	printOptions(file_names)
	print('')
	
	file_names = createDocFileNames(file_names)

	choice = chooseFile()
	print('')

	if not (choice >= 1 and choice <= len(file_names)):
		print("Choice is out of range. Exiting program.")
		exit(1)

	# print("File names:")
	# print(file_names)
	file_name = file_names[choice-1]
	print("%s selected\n"%file_name)

	createDoc(file_name)
	print('')

	addTask(file_name)


if __name__ == '__main__':
	main()
